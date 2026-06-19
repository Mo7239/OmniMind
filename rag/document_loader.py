import os
from pathlib import Path
from typing import List
from core.logger import get_logger

logger = get_logger(__name__)

class Document:
    def __init__(self, content : str, metadata : dict ={}):
        self.content = content
        self.metadata = metadata

    def __repr__(self):
        return f"Document(chars={len(self.content)}, metadata={self.metadata})"

class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".txt", ".pdf", ".docx"}   

    def load(self, path:str) -> List[Document] :
        file_path = Path(path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {path}")
        
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        logger.info(f"Loading file | path={path} | type={ext}")

        if ext ==".txt":
            return self._load_txt(file_path)
        elif ext ==".pdf":
            return self._load_pdf(file_path)
        elif ext==".docx":
            return self._load_docx(file_path)
        
    def load_directory(self, dir_path:str) -> List[Document]:
        directory = Path(dir_path)
        if not directory.exists():
            raise FileNotFoundError(f"Directory not found: {dir_path}")

        docs = []
        for file in directory.iterdir():
            if file.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                docs.extend(self.load(str(file)))


        logger.info(f"Loaded {len(docs)} documents from {dir_path}")
        return docs

    def _load_txt(self, path:Path) -> List[Document]:
        with open(path,"r",encoding="utf-8") as f:
            content = f.read()
        return [Document(content=content,metadata={"source": str(path), "type": "txt"})]
        
    def _load_pdf(self,path:Path) -> List[Document]:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        docs = []
        for idx , page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                docs.append(Document(
                    content=text,
                    metadata={"source": str(path), "type": "pdf","page":idx + 1}
                ))
        return docs  

    def _load_docx(self, path : Path) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(str(path))
        content = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        return [Document(content=content, metadata={"source": str(path), "type": "docx"})]
    


         
    
        
